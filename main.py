from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Cookie, status
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
import pandas as pd
import os
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import shutil
import subprocess
import uuid
import re
import zipfile
import json
import hashlib
import logging
import secrets
import datetime
from io import BytesIO
from typing import Dict, Any, List, Set, Optional
from pypdf import PdfReader, PdfWriter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()
USERS_FILE = "users.json"
TICKETS_FILE = "tickets.json"
SESSION_COOKIE_NAME = "session_token"
SESSION_STORE: Dict[str, dict] = {}

PERMISSIONS = {
    "viewer": {"visualizar"},
    "editor": {"visualizar", "editar"},
    "creator": {"visualizar", "editar", "crear"},
}

DEFAULT_USERS = [
    {"username": "viewer", "password": "viewer123", "role": "viewer"},
    {"username": "editor", "password": "editor123", "role": "editor"},
    {"username": "creator", "password": "creator123", "role": "creator"},
]

# Crear directorios necesarios si no existen
os.makedirs("static", exist_ok=True)
os.makedirs("output", exist_ok=True)
os.makedirs("output/certificados", exist_ok=True)
os.makedirs("output/tickets", exist_ok=True)
os.makedirs("output/previews", exist_ok=True)
os.makedirs("uploads", exist_ok=True)
os.makedirs("templates", exist_ok=True)

app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/output", StaticFiles(directory="output"), name="output")
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")
app.mount("/ticket-assets", StaticFiles(directory="output/tickets"), name="ticket-assets")

# Inicializar templates aquí, antes de usar en las rutas
templates = Jinja2Templates(directory="templates")


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()


def save_users(users):
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, indent=2, ensure_ascii=False)


def load_users():
    if not os.path.exists(USERS_FILE):
        usr_copy = []
        for u in DEFAULT_USERS:
            usr_copy.append({"username": u["username"], "password": hash_password(u["password"]), "role": u["role"]})
        save_users(usr_copy)

    with open(USERS_FILE, "r", encoding="utf-8") as f:
        datos = json.load(f)
    return datos


def get_user(username: str):
    for u in load_users():
        if u["username"] == username:
            return u
    return None


def load_tickets() -> List[Dict[str, Any]]:
    if not os.path.exists(TICKETS_FILE):
        with open(TICKETS_FILE, "w", encoding="utf-8") as f:
            json.dump([], f, indent=2, ensure_ascii=False)
        return []

    with open(TICKETS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def save_tickets(tickets: List[Dict[str, Any]]) -> None:
    with open(TICKETS_FILE, "w", encoding="utf-8") as f:
        json.dump(tickets, f, indent=2, ensure_ascii=False)


def get_ticket(ticket_id: str) -> Optional[Dict[str, Any]]:
    tickets = load_tickets()
    for ticket in tickets:
        if ticket.get("id") == ticket_id:
            return ticket
    return None


def authenticate_user(username: str, password: str):
    user = get_user(username)
    if not user:
        return None
    if user["password"] != hash_password(password):
        return None
    return user


def get_current_user(session_token: Optional[str] = Cookie(None)):
    if not session_token or session_token not in SESSION_STORE:
        logger.warning("No session token or invalid token")
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="No autenticado")

    session_data = SESSION_STORE[session_token]
    username = session_data["username"]
    user = get_user(username)
    if not user:
        logger.error(f"User {username} not found in users database")
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Usuario inválido")
    return user


def require_permission(permission: str):
    def permission_dependency(user: dict = Depends(get_current_user)):
        role = user.get("role")
        if permission not in PERMISSIONS.get(role, set()):
            raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Permisos insuficientes")
        return user

    return permission_dependency


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    return templates.TemplateResponse(
        request,
        "login.html",
        {"request": request, "message": "Inicia sesión"},
    )


@app.post("/login")
async def login(request: Request, username: str = Form(...), password: str = Form(...)):
    logger.info(f"Login attempt for user: {username}")
    user = authenticate_user(username, password)
    if not user:
        logger.warning(f"Failed login attempt for user: {username}")
        return templates.TemplateResponse(
            request,
            "login.html",
            {"request": request, "message": "Credenciales incorrectas"},
        )

    token = secrets.token_hex(32)
    SESSION_STORE[token] = {
        "username": username,
        "created_at": secrets.token_hex(16)  # Simple timestamp
    }
    logger.info(f"Successful login for user: {username}, token: {token[:8]}...")

    response = RedirectResponse(url="/dashboard", status_code=status.HTTP_303_SEE_OTHER)
    response.set_cookie(
        SESSION_COOKIE_NAME,
        token,
        httponly=True,
        samesite="lax",
        secure=True,  # Importante para HTTPS en Render
        max_age=3600  # 1 hora
    )
    return response


@app.get("/logout")
def logout(session_token: Optional[str] = Cookie(None)):
    if session_token and session_token in SESSION_STORE:
        del SESSION_STORE[session_token]
        logger.info(f"User logged out, session {session_token[:8]}... cleared")
    response = RedirectResponse(url="/login", status_code=status.HTTP_303_SEE_OTHER)
    response.delete_cookie(SESSION_COOKIE_NAME)
    return response


@app.get("/", response_class=HTMLResponse)
def home(request: Request, session_token: Optional[str] = Cookie(None)):
    # Si tiene sesión válida, ir a dashboard; si no, ir a login
    if session_token and session_token in SESSION_STORE:
        return RedirectResponse(url="/dashboard")
    else:
        return RedirectResponse(url="/login")


@app.get("/dashboard", response_class=HTMLResponse)
def dashboard(request: Request, user: dict = Depends(get_current_user)):
    try:
        logger.info(f"Dashboard access for user: {user['username']}")
        can_preview = "editar" in PERMISSIONS.get(user.get("role"), set()) or "crear" in PERMISSIONS.get(user.get("role"), set())
        can_generate = "crear" in PERMISSIONS.get(user.get("role"), set())
        ticket_history = []
        if user.get("role") == "creator":
            tickets = load_tickets()
            ticket_history = [
                {
                    "id": t.get("id"),
                    "created_at": t.get("created_at"),
                    "row_count": len(t.get("rows", [])),
                    "plantilla_nombre": t.get("plantilla_nombre"),
                }
                for t in tickets
                if t.get("created_by") == user.get("username")
            ]
        return templates.TemplateResponse(
            request,
            "formulario.html",
            {
                "request": request,
                "user": user,
                "can_preview": can_preview,
                "can_generate": can_generate,
                "message": "Bienvenido",
                "ticket_history": ticket_history,
            },
        )
    except Exception as e:
        logger.error(f"Error in dashboard for user {user.get('username', 'unknown')}: {e}")
        # En caso de error, redirigir al login
        return RedirectResponse(url="/login", status_code=status.HTTP_303_SEE_OTHER)


def create_ticket_preview_session(ticket: Dict[str, Any]) -> str:
    session_id = uuid.uuid4().hex
    ticket_folder = os.path.join("output", "tickets", ticket["id"])

    imagenes = []
    if ticket.get("imagenes"):
        for item in ticket.get("imagenes", []):
            imagenes.append(
                {
                    "filename": item.get("filename"),
                    "original_name": item.get("original_name", item.get("filename")),
                    "path": os.path.join(ticket_folder, item.get("filename")),
                    "x": item.get("x", 0.0),
                    "y": item.get("y", 0.0),
                    "width": item.get("width", DEFAULT_SELLO_WIDTH_IN),
                    "page": item.get("page", 1),
                }
            )
    else:
        if ticket.get("sello_nombre"):
            imagenes.append(
                {
                    "filename": ticket.get("sello_nombre"),
                    "original_name": ticket.get("sello_nombre"),
                    "path": os.path.join(ticket_folder, ticket.get("sello_nombre")),
                    "x": ticket.get("sello_x", 0.0),
                    "y": ticket.get("sello_y", 0.0),
                    "width": ticket.get("sello_width_in", DEFAULT_SELLO_WIDTH_IN),
                    "page": ticket.get("sello_page", 1),
                }
            )
        if ticket.get("firma_nombre"):
            imagenes.append(
                {
                    "filename": ticket.get("firma_nombre"),
                    "original_name": ticket.get("firma_nombre"),
                    "path": os.path.join(ticket_folder, ticket.get("firma_nombre")),
                    "x": ticket.get("firma_x", 0.0),
                    "y": ticket.get("firma_y", 0.0),
                    "width": ticket.get("firma_width_in", DEFAULT_FIRMA_WIDTH_IN),
                    "page": ticket.get("firma_page", 1),
                }
            )

    PREVIEW_SESSIONS[session_id] = {
        "ticket_id": ticket["id"],
        "plantilla_path": os.path.join(ticket_folder, ticket["plantilla_nombre"]),
        "plantilla_nombre": ticket["plantilla_nombre"],
        "rows": ticket.get("rows", []),
        "variables": ticket.get("variables", []),
        "imagenes": imagenes,
        "page_width_in": ticket.get("page_width_in", 8.27),
        "page_height_in": ticket.get("page_height_in", 11.69),
    }
    return session_id


def save_ticket_assets(session: Dict[str, Any], ticket_id: str) -> str:
    ticket_folder = os.path.join("output", "tickets", ticket_id)
    os.makedirs(ticket_folder, exist_ok=True)

    plantilla_destino = os.path.join(ticket_folder, session.get("plantilla_nombre"))
    if session.get("plantilla_path") and os.path.exists(session.get("plantilla_path")):
        shutil.copy2(session.get("plantilla_path"), plantilla_destino)

    for imagen in session.get("imagenes", []):
        if imagen.get("path") and os.path.exists(imagen.get("path")):
            destino = os.path.join(ticket_folder, imagen.get("filename"))
            shutil.copy2(imagen.get("path"), destino)

    return ticket_folder


@app.get("/tickets", response_class=HTMLResponse)
def listar_tickets(request: Request, user: dict = Depends(require_permission("crear"))):
    tickets = [
        t for t in load_tickets() if t.get("created_by") == user.get("username")
    ]
    tickets.sort(key=lambda t: t.get("created_at", ""), reverse=True)
    return templates.TemplateResponse(
        request,
        "tickets.html",
        {
            "request": request,
            "user": user,
            "tickets": tickets,
        },
    )


@app.get("/tickets/{ticket_id}", response_class=HTMLResponse)
def detalle_ticket(ticket_id: str, request: Request, user: dict = Depends(require_permission("crear"))):
    ticket = get_ticket(ticket_id)
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket no encontrado")
    if ticket.get("created_by") != user.get("username"):
        raise HTTPException(status_code=403, detail="No tienes permiso para ver este ticket")

    return templates.TemplateResponse(
        request,
        "ticket_detalle.html",
        {
            "request": request,
            "user": user,
            "ticket": ticket,
        },
    )


@app.get("/tickets/{ticket_id}/editar", response_class=HTMLResponse)
def editar_ticket(ticket_id: str, request: Request, user: dict = Depends(require_permission("crear"))):
    ticket = get_ticket(ticket_id)
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket no encontrado")
    if ticket.get("created_by") != user.get("username"):
        raise HTTPException(status_code=403, detail="No tienes permiso para editar este ticket")

    session_id = create_ticket_preview_session(ticket)
    session = PREVIEW_SESSIONS[session_id]

    return templates.TemplateResponse(
        request,
        "previsualizacion.html",
        {
            "request": request,
            "user": user,
            "session_id": session_id,
            "rows": list(enumerate(session["rows"])),
            "variables": sorted(session["variables"]),
            "variables_json": json.dumps(sorted(session["variables"])),
            "plantilla_nombre": session.get("plantilla_nombre"),
            "imagenes": session.get("imagenes", []),
            "imagenes_json": json.dumps(session.get("imagenes", [])),
            "page_width_in": session.get("page_width_in", 8.27),
            "page_height_in": session.get("page_height_in", 11.69),
            "ticket_id": ticket_id,
            "editing_ticket": True,
        },
    )


@app.get("/tickets/{ticket_id}/editar/{row_id}")
def editar_ticket_row(ticket_id: str, row_id: int, user: dict = Depends(require_permission("crear"))):
    ticket = get_ticket(ticket_id)
    if not ticket:
        raise HTTPException(status_code=404, detail="Ticket no encontrado")
    if ticket.get("created_by") != user.get("username"):
        raise HTTPException(status_code=403, detail="No tienes permiso para editar este ticket")

    if row_id < 0 or row_id >= len(ticket.get("rows", [])):
        raise HTTPException(status_code=404, detail="Fila no encontrada")

    session_id = create_ticket_preview_session(ticket)
    return RedirectResponse(url=f"/visor-preview/{session_id}/{row_id}")


@app.get("/certificados", response_class=HTMLResponse)
def listar_certificados(request: Request, user: dict = Depends(require_permission("visualizar"))):
    carpeta = "output/certificados"
    archivos = []
    if os.path.exists(carpeta):
        archivos = [f for f in os.listdir(carpeta) if os.path.isfile(os.path.join(carpeta, f))]
        archivos.sort()

    return templates.TemplateResponse(
        request,
        "resultado.html",
        {
            "request": request,
            "generados": [{"nombre": a, "pdf": f"/output/certificados/{a}"} for a in archivos],
            "user": user,
            "message": "Lista de certificados disponibles",
        },
    )


@app.get("/certificados/{archivo}")
def descargar_certificado(archivo: str, user: dict = Depends(require_permission("visualizar"))):
    ruta = os.path.join("output/certificados", archivo)
    if not os.path.exists(ruta):
        raise HTTPException(status_code=404, detail="Certificado no encontrado")
    return FileResponse(path=ruta, filename=archivo, media_type="application/pdf")

# Crear directorios antes de montarlos como archivos estáticos
os.makedirs("output/certificados", exist_ok=True)
os.makedirs("output/previews", exist_ok=True)
os.makedirs("uploads", exist_ok=True)

PREVIEW_SESSIONS: Dict[str, Dict[str, Any]] = {}
DEFAULT_SELLO_WIDTH_IN = 2.5
DEFAULT_FIRMA_WIDTH_IN = 3.0


def convertir_docx_a_pdf(ruta_docx: str, carpeta_salida: str) -> str:
    """
    Convierte un DOCX a PDF usando LibreOffice (soffice).
    Devuelve la ruta del PDF generado.
    """
    exe = os.getenv("SOFFICE_PATH", "soffice")
    comando = [
        exe,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        carpeta_salida,
        ruta_docx,
    ]

    try:
        subprocess.run(comando, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except FileNotFoundError:
        raise RuntimeError(
            "No se encontro LibreOffice. Define SOFFICE_PATH o agrega 'soffice' al PATH."
        )
    except subprocess.CalledProcessError as exc:
        salida = exc.stderr.decode(errors="ignore")
        raise RuntimeError(f"Error al convertir a PDF: {salida}")

    base = os.path.splitext(os.path.basename(ruta_docx))[0]
    return os.path.join(carpeta_salida, f"{base}.pdf")


def overlay_imagenes_en_pdf(
    ruta_pdf_entrada: str,
    ruta_pdf_salida: str,
    overlays: List[Dict[str, Any]],
) -> None:
    """
    Dibuja imagenes por pagina sobre un PDF ya generado.
    Coordenadas en pulgadas desde esquina superior-izquierda.
    """
    reader = PdfReader(ruta_pdf_entrada)
    writer = PdfWriter()

    overlays_por_pagina: Dict[int, List[Dict[str, Any]]] = {}
    for ov in overlays:
        idx = int(max(1, ov.get("page", 1))) - 1
        overlays_por_pagina.setdefault(idx, []).append(ov)

    for page_idx, page in enumerate(reader.pages):
        if page_idx in overlays_por_pagina:
            ancho_pt = float(page.mediabox.width)
            alto_pt = float(page.mediabox.height)
            packet = BytesIO()
            c = canvas.Canvas(packet, pagesize=(ancho_pt, alto_pt))

            for ov in overlays_por_pagina[page_idx]:
                image_path = ov.get("path")
                if not image_path or not os.path.exists(image_path):
                    continue

                x_in = float(max(0.0, ov.get("x_in", 0.0)))
                y_in = float(max(0.0, ov.get("y_in", 0.0)))
                width_in = float(max(0.1, ov.get("width_in", 1.0)))

                img_reader = ImageReader(image_path)
                img_w_px, img_h_px = img_reader.getSize()
                aspect = (img_h_px / img_w_px) if img_w_px else 1.0
                width_pt = width_in * 72.0
                height_pt = width_pt * aspect
                x_pt = x_in * 72.0
                y_pt = alto_pt - (y_in * 72.0) - height_pt

                c.drawImage(
                    img_reader,
                    x_pt,
                    y_pt,
                    width=width_pt,
                    height=height_pt,
                    preserveAspectRatio=True,
                    mask="auto",
                )

            c.save()
            packet.seek(0)
            overlay_page = PdfReader(packet).pages[0]
            page.merge_page(overlay_page)

        writer.add_page(page)

    with open(ruta_pdf_salida, "wb") as out_file:
        writer.write(out_file)


def limpiar_nombre_archivo(nombre: str) -> str:
    limpio = re.sub(r'[\\/:*?"<>|]+', "_", str(nombre)).strip()
    return limpio or "certificado"


def obtener_tamano_pagina_pulgadas(ruta_docx: str) -> Dict[str, float]:
    """Lee tamaño real de página (en pulgadas) desde la primera sección del DOCX."""
    doc = Document(ruta_docx)
    seccion = doc.sections[0]
    return {
        "page_width_in": float(seccion.page_width.inches),
        "page_height_in": float(seccion.page_height.inches),
    }


def agregar_imagen_flotante_absoluta(
    doc: Document,
    image_path: str,
    x_in: float,
    y_in: float,
    width_in: float,
) -> None:
    """
    Inserta imagen como objeto flotante con posicion absoluta (wp:anchor),
    relativo al borde superior-izquierdo de la página.
    """
    parrafo = doc.add_paragraph()
    run = parrafo.add_run()
    inline_shape = run.add_picture(image_path, width=Inches(width_in))
    inline = inline_shape._inline

    cx = inline.extent.cx
    cy = inline.extent.cy
    doc_pr = inline.docPr
    doc_pr_id = int(doc_pr.get("id"))
    doc_pr_name = doc_pr.get("name", f"Picture {doc_pr_id}")
    graphic_xml = inline.graphic.xml
    x_emu = int(max(0, x_in) * 914400)
    y_emu = int(max(0, y_in) * 914400)

    anchor = parse_xml(
        f"""
        <wp:anchor distT="0" distB="0" distL="0" distR="0"
            simplePos="0" relativeHeight="251658240" behindDoc="0"
            locked="0" layoutInCell="1" allowOverlap="1"
            {nsdecls('wp', 'a', 'pic', 'r')}>
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page">
                <wp:posOffset>{x_emu}</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="page">
                <wp:posOffset>{y_emu}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{cx}" cy="{cy}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="{doc_pr_id}" name="{doc_pr_name}"/>
            <wp:cNvGraphicFramePr>
                <a:graphicFrameLocks noChangeAspect="1"/>
            </wp:cNvGraphicFramePr>
            {graphic_xml}
        </wp:anchor>
        """
    )
    inline.getparent().replace(inline, anchor)


def aplicar_marca_y_elementos(
    ruta_docx: str,
    no_valido: bool = False,
    sello_path: str = None,
    firma_path: str = None,
    sello_x: float = 0.0,
    sello_y: float = 0.0,
    firma_x: float = 0.0,
    firma_y: float = 0.0,
    incluir_imagenes: bool = False,
    sello_width_in: float = DEFAULT_SELLO_WIDTH_IN,
    firma_width_in: float = DEFAULT_FIRMA_WIDTH_IN,
) -> None:
    """Agrega marca de agua textual + imágenes de sello/firma al DOCX final."""
    try:
        doc = Document(ruta_docx)

        if no_valido:
            for section in doc.sections:
                header = section.header
                if not header.paragraphs:
                    par = header.add_paragraph()
                else:
                    par = header.paragraphs[0]
                if hasattr(par, "clear"):
                    par.clear()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = par.add_run("NO VÁLIDO")
                run.font.size = Pt(48)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)

        if incluir_imagenes and sello_path and os.path.exists(sello_path):
            agregar_imagen_flotante_absoluta(
                doc=doc,
                image_path=sello_path,
                x_in=sello_x,
                y_in=sello_y,
                width_in=sello_width_in,
            )

        if incluir_imagenes and firma_path and os.path.exists(firma_path):
            agregar_imagen_flotante_absoluta(
                doc=doc,
                image_path=firma_path,
                x_in=firma_x,
                y_in=firma_y,
                width_in=firma_width_in,
            )

        doc.save(ruta_docx)
    except Exception as e:
        print(f"Error aplicando marca o imagenes: {e}")


def extraer_variables_plantilla(ruta_plantilla: str) -> Set[str]:
    """
    Extrae todas las variables {{variable}} de una plantilla DOCX.
    Lee el XML interno del DOCX (que es un ZIP) y busca patrones.
    """
    variables = set()
    try:
        with zipfile.ZipFile(ruta_plantilla, 'r') as docx:
            # Los documentos DOCX tienen el contenido en document.xml
            with docx.open('word/document.xml') as xml_file:
                contenido = xml_file.read().decode('utf-8')
                # Buscar patrones {{ variable }}
                matches = re.findall(r'\{\{\s*(\w+)\s*\}\}', contenido)
                variables.update(matches)
    except Exception as e:
        print(f"Error extrayendo variables: {e}")
    
    return variables


def cargar_registros_excel(file_obj, variables_requeridas: Set[str]) -> List[Dict[str, str]]:
    """
    Carga registros desde Excel detectando dinámicamente las columnas necesarias.
    """
    df = pd.read_excel(file_obj)
    
    # Hacer case-insensitive: crear mapeo de columnas en minúsculas
    df_columns_lower = {col.lower(): col for col in df.columns}
    
    # Buscar columnas (case-insensitive)
    faltantes = []
    columnas_mapeo = {}
    
    for var in variables_requeridas:
        var_lower = var.lower()
        if var_lower in df_columns_lower:
            columnas_mapeo[var] = df_columns_lower[var_lower]
        else:
            faltantes.append(var)
    
    if faltantes:
        raise HTTPException(
            status_code=400,
            detail=f"Faltan columnas en Excel: {', '.join(sorted(faltantes))}. Se esperaban: {', '.join(sorted(variables_requeridas))}",
        )

    registros: List[Dict[str, str]] = []
    for _, row in df.iterrows():
        registro = {}
        tiene_datos = False
        
        for var, col_real in columnas_mapeo.items():
            valor = "" if pd.isna(row[col_real]) else str(row[col_real]).strip()
            registro[var] = valor
            if valor:
                tiene_datos = True
        
        if tiene_datos:
            registros.append(registro)

    return registros


def render_docx_desde_datos(
    ruta_plantilla: str,
    datos: Dict[str, str],
    ruta_docx: str,
    no_valido: bool,
    sello_path: str = None,
    firma_path: str = None,
    sello_x: float = 0.0,
    sello_y: float = 0.0,
    firma_x: float = 0.0,
    firma_y: float = 0.0,
    incluir_imagenes_docx: bool = False,
    sello_width_in: float = DEFAULT_SELLO_WIDTH_IN,
    firma_width_in: float = DEFAULT_FIRMA_WIDTH_IN,
) -> None:
    """
    Renderiza DOCX con datos dinámicos.
    datos debe ser un diccionario con las variables de la plantilla.
    """
    doc = DocxTemplate(ruta_plantilla)
    
    # Preparar contexto: agregar variables dinámicas + no_valido
    contexto = dict(datos)
    contexto["no_valido"] = no_valido
    if no_valido:
        contexto["watermark"] = "NO VALIDO, DOCUMENTO NO OFICIAL"
    
    doc.render(contexto)
    doc.save(ruta_docx)

    # Agregar marca en encabezado y sellos/firma si existen
    aplicar_marca_y_elementos(
        ruta_docx,
        no_valido=no_valido,
        sello_path=sello_path,
        firma_path=firma_path,
        sello_x=sello_x,
        sello_y=sello_y,
        firma_x=firma_x,
        firma_y=firma_y,
        incluir_imagenes=incluir_imagenes_docx,
        sello_width_in=sello_width_in,
        firma_width_in=firma_width_in,
    )





@app.post("/previsualizar", response_class=HTMLResponse)
async def previsualizar(
    request: Request,
    file: UploadFile = File(...),
    plantilla: UploadFile = File(...),
    imagenes: List[UploadFile] = File(None),
    sello_x: float = Form(0.0),
    sello_y: float = Form(0.0),
    firma_x: float = Form(0.0),
    firma_y: float = Form(0.0),
    user: dict = Depends(require_permission("editar")),
):
    if not plantilla.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="La plantilla debe ser un archivo .docx")

    # Guardar plantilla temporalmente para extraer variables
    session_id = uuid.uuid4().hex
    plantilla_id = f"{session_id}_{os.path.basename(plantilla.filename)}"
    ruta_plantilla = os.path.join("uploads", plantilla_id)

    with open(ruta_plantilla, "wb") as f:
        await plantilla.seek(0)
        shutil.copyfileobj(plantilla.file, f)

    # Extraer variables de la plantilla
    variables = extraer_variables_plantilla(ruta_plantilla)
    if not variables:
        raise HTTPException(status_code=400, detail="No se encontraron variables en la plantilla DOCX")

    # Cargar Excel con las variables detectadas
    registros = cargar_registros_excel(file.file, variables)
    if not registros:
        raise HTTPException(status_code=400, detail="No hay registros validos para previsualizar")

    imagenes_cargadas = []
    if imagenes:
        for idx_img, imagen in enumerate(imagenes):
            if imagen and imagen.filename:
                imagen_id = f"{session_id}_img_{idx_img}_{os.path.basename(imagen.filename)}"
                imagen_path = os.path.join("uploads", imagen_id)
                with open(imagen_path, "wb") as f:
                    await imagen.seek(0)
                    shutil.copyfileobj(imagen.file, f)
                imagenes_cargadas.append(
                    {
                        "filename": os.path.basename(imagen_id),
                        "original_name": os.path.basename(imagen.filename),
                        "path": imagen_path,
                        "x": sello_x,
                        "y": sello_y,
                        "width": DEFAULT_SELLO_WIDTH_IN,
                        "page": 1,
                    }
                )

    PREVIEW_SESSIONS[session_id] = {
        "plantilla_path": ruta_plantilla,
        "plantilla_nombre": os.path.basename(plantilla.filename),
        "rows": registros,
        "variables": variables,
        "imagenes": imagenes_cargadas,
        "sello_x": sello_x,
        "sello_y": sello_y,
        "firma_x": firma_x,
        "firma_y": firma_y,
        **obtener_tamano_pagina_pulgadas(ruta_plantilla),
    }

    return templates.TemplateResponse(
        request,
        "previsualizacion.html",
        {
            "request": request,
            "user": user,
            "session_id": session_id,
            "rows": list(enumerate(registros)),
            "variables": sorted(variables),
            "variables_json": json.dumps(sorted(variables)),
            "plantilla_nombre": os.path.basename(plantilla.filename),
            "imagenes": imagenes_cargadas,
            "imagenes_json": json.dumps(imagenes_cargadas),
            "ticket_id": None,
            "page_width_in": PREVIEW_SESSIONS[session_id]["page_width_in"],
            "page_height_in": PREVIEW_SESSIONS[session_id]["page_height_in"],
        },
    )


@app.post("/session/subir-imagenes")
async def subir_imagenes_sesion(
    request: Request,
    session_id: str = Form(...),
    imagenes: List[UploadFile] = File(None),
    user: dict = Depends(require_permission("editar")),
):
    session = PREVIEW_SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Sesion de previsualizacion no encontrada")

    if not imagenes:
        return JSONResponse({"ok": False, "message": "Selecciona una o más imágenes para subir."}, status_code=400)

    existing_count = len(session.get("imagenes", []))
    uploaded = []
    for idx, imagen in enumerate(imagenes):
        if imagen and imagen.filename:
            imagen_id = f"{session_id}_img_{existing_count + idx}_{os.path.basename(imagen.filename)}"
            imagen_path = os.path.join("uploads", imagen_id)
            with open(imagen_path, "wb") as f:
                await imagen.seek(0)
                shutil.copyfileobj(imagen.file, f)
            new_item = {
                "filename": os.path.basename(imagen_id),
                "original_name": os.path.basename(imagen.filename),
                "path": imagen_path,
                "x": 0.0,
                "y": 0.0,
                "width": DEFAULT_SELLO_WIDTH_IN,
                "page": 1,
            }
            session.setdefault("imagenes", []).append(new_item)
            uploaded.append(new_item)

    if not uploaded:
        return JSONResponse({"ok": False, "message": "No se subieron imágenes válidas."}, status_code=400)

    image_response = []
    ticket_asset_dir = None
    if session.get("ticket_id"):
        ticket_asset_dir = os.path.abspath(os.path.join("output", "tickets", session["ticket_id"]))

    for item in session.get("imagenes", []):
        item_path = os.path.abspath(item.get("path", "")) if item.get("path") else ""
        if ticket_asset_dir and item_path.startswith(ticket_asset_dir):
            src = f"/ticket-assets/{session['ticket_id']}/{item.get('filename')}"
        else:
            src = f"/uploads/{item.get('filename')}"
        image_response.append({
            "filename": item.get("filename"),
            "original_name": item.get("original_name"),
            "x": item.get("x", 0.0),
            "y": item.get("y", 0.0),
            "width": item.get("width", DEFAULT_SELLO_WIDTH_IN),
            "page": item.get("page", 1),
            "src": src,
        })

    return JSONResponse({"ok": True, "message": "Imágenes agregadas.", "imagenes": image_response})


@app.get("/preview/{session_id}/{row_id}")
async def preview_pdf(request: Request, session_id: str, row_id: int):
    session = PREVIEW_SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Sesion de previsualizacion no encontrada")

    if row_id < 0 or row_id >= len(session["rows"]):
        raise HTTPException(status_code=404, detail="Registro no encontrado")

    # Obtener datos del registro
    registro_base = session["rows"][row_id]
    
    # Obtener datos editados del formulario si existen
    query_params = request.query_params
    datos = {}
    for var in session["variables"]:
        # Buscar en query params o usar el valor base
        datos[var] = query_params.get(var, registro_base.get(var, ""))

    carpeta_preview = os.path.join("output", "previews", session_id)
    os.makedirs(carpeta_preview, exist_ok=True)

    # Usar primera variable para nombrar archivo
    primera_var = list(session["variables"])[0] if session["variables"] else "preview"
    base_nombre = limpiar_nombre_archivo(f"preview_{row_id}_{datos.get(primera_var, '')}")
    ruta_docx = os.path.join(carpeta_preview, f"{base_nombre}.docx")

    try:
        render_docx_desde_datos(
            session["plantilla_path"],
            datos=datos,
            ruta_docx=ruta_docx,
            no_valido=True,
            sello_path=session.get("sello_path"),
            firma_path=session.get("firma_path"),
            sello_x=session.get("sello_x", 0.0),
            sello_y=session.get("sello_y", 0.0),
            firma_x=session.get("firma_x", 0.0),
            firma_y=session.get("firma_y", 0.0),
            incluir_imagenes_docx=False,
            sello_width_in=session.get("sello_width_in", DEFAULT_SELLO_WIDTH_IN),
            firma_width_in=session.get("firma_width_in", DEFAULT_FIRMA_WIDTH_IN),
        )
        ruta_pdf_base = convertir_docx_a_pdf(ruta_docx, carpeta_preview)
        ruta_pdf = ruta_pdf_base.replace(".pdf", "_overlay.pdf")
        overlays = []
        
        # Usar posiciones individuales si existen, si no usar globales
        rows_overrides = session.get("rows_overrides", {})
        imagenes_source = session.get("imagenes", [])
        if str(row_id) in rows_overrides:
            imagenes_source = rows_overrides[str(row_id)]
        
        for imagen in imagenes_source:
            overlays.append(
                {
                    "path": imagen.get("path"),
                    "x_in": imagen.get("x", 0.0),
                    "y_in": imagen.get("y", 0.0),
                    "width_in": imagen.get("width", DEFAULT_SELLO_WIDTH_IN),
                    "page": imagen.get("page", 1),
                }
            )
        overlay_imagenes_en_pdf(
            ruta_pdf_entrada=ruta_pdf_base,
            ruta_pdf_salida=ruta_pdf,
            overlays=overlays,
        )
    except (RuntimeError, Exception) as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    filename = os.path.basename(ruta_pdf)
    return FileResponse(
        path=ruta_pdf,
        media_type="application/pdf",
        filename=filename,
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )


@app.get("/visor-preview/{session_id}/{row_id}", response_class=HTMLResponse)
def visor_preview(request: Request, session_id: str, row_id: int, user: dict = Depends(require_permission("visualizar"))):
    session = PREVIEW_SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Sesion no encontrada")
    
    # Obtener datos editados desde query params
    query_params = request.query_params
    datos = {}
    for var in session["variables"]:
        datos[var] = query_params.get(var, session["rows"][row_id].get(var, ""))
    
    # Usar posiciones individuales si existen, si no usar globales
    imagenes_source = session.get("imagenes", [])
    rows_overrides = session.get("rows_overrides", {})
    if str(row_id) in rows_overrides:
        imagenes_source = rows_overrides[str(row_id)]
    
    imagenes = []
    for item in imagenes_source:
        imagenes.append(
            {
                "filename": item.get("filename"),
                "original_name": item.get("original_name", item.get("filename")),
                "src": f"/ticket-assets/{session['ticket_id']}/{item['filename']}" if session.get("ticket_id") else f"/uploads/{item.get('filename')}",
                "x": item.get("x", 0.0),
                "y": item.get("y", 0.0),
                "width": item.get("width", DEFAULT_SELLO_WIDTH_IN),
                "page": item.get("page", 1),
            }
        )

    return templates.TemplateResponse(
        request,
        "visor_preview.html",
        {
            "request": request,
            "user": user,
            "session_id": session_id,
            "row_id": row_id,
            "plantilla_nombre": session.get("plantilla_nombre"),
            "datos": datos,
            "datos_json": json.dumps(datos),
            "variables": sorted(session["variables"]),
            "imagenes": imagenes,
            "imagenes_json": json.dumps(imagenes),
            "ticket_id": session.get("ticket_id"),
            "page_width_in": session.get("page_width_in", 8.27),
            "page_height_in": session.get("page_height_in", 11.69),
        },
    )


@app.post("/session/ajustar-posicion")
def ajustar_posicion(
    session_id: str = Form(...),
    imagenes_json: str = Form("[]"),
    edit_scope: str = Form("individual"),
    row_id: int = Form(0),
    user: dict = Depends(require_permission("editar")),
):
    session = PREVIEW_SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Sesion no encontrada")

    try:
        imagenes = json.loads(imagenes_json)
        if isinstance(imagenes, list):
            actuales = {img.get("filename"): img for img in session.get("imagenes", [])}
            imagenes_procesadas = []
            for item in imagenes:
                if not item.get("filename"):
                    continue
                filename = item.get("filename")
                existing = actuales.get(filename, {})
                imagenes_procesadas.append(
                    {
                        "filename": filename,
                        "original_name": item.get("original_name") or existing.get("original_name"),
                        "path": existing.get("path"),
                        "x": float(item.get("x", 0.0)),
                        "y": float(item.get("y", 0.0)),
                        "page": int(item.get("page", 1)),
                        "width": float(item.get("width", DEFAULT_SELLO_WIDTH_IN)),
                    }
                )
            
            # Si es edición global, actualizar posiciones para todas las filas
            if edit_scope == "global":
                session["imagenes"] = imagenes_procesadas
                message = "Posiciones actualizadas globalmente"
            # Si es edición individual, guardar override solo para esa fila
            elif edit_scope == "individual":
                # Validar que row_id es válido
                if row_id < 0 or row_id >= len(session.get("rows", [])):
                    raise HTTPException(status_code=400, detail="Fila no válida")
                # Guardar posiciones específicas para esa fila
                if "rows_overrides" not in session:
                    session["rows_overrides"] = {}
                session["rows_overrides"][str(row_id)] = imagenes_procesadas
                message = f"Posiciones actualizadas para el certificado #{row_id + 1}"
            else:
                message = "Posiciones actualizadas"
            
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Formato JSON invalido para las imagenes")

    return {"ok": True, "message": message}


@app.post("/generar-final", response_class=HTMLResponse)
async def generar_final(
    request: Request,
    session_id: str = Form(...),
    selected_ids: List[str] = Form(default=[]),
    user: dict = Depends(require_permission("crear")),
):
    session = PREVIEW_SESSIONS.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Sesion de previsualizacion no encontrada")

    form = await request.form()
    ticket_id = form.get("ticket_id")

    ids_seleccionados = selected_ids
    if isinstance(ids_seleccionados, str):
        ids_seleccionados = [ids_seleccionados]

    if not ids_seleccionados:
        raise HTTPException(status_code=400, detail="Debes seleccionar al menos un certificado")

    sello_x = float(form.get("sello_x", session.get("sello_x", 0.0)) or 0.0)
    sello_y = float(form.get("sello_y", session.get("sello_y", 0.0)) or 0.0)
    firma_x = float(form.get("firma_x", session.get("firma_x", 0.0)) or 0.0)
    firma_y = float(form.get("firma_y", session.get("firma_y", 0.0)) or 0.0)
    sello_width_in = float(form.get("sello_width_in", session.get("sello_width_in", DEFAULT_SELLO_WIDTH_IN)) or DEFAULT_SELLO_WIDTH_IN)
    firma_width_in = float(form.get("firma_width_in", session.get("firma_width_in", DEFAULT_FIRMA_WIDTH_IN)) or DEFAULT_FIRMA_WIDTH_IN)
    sello_page = int(form.get("sello_page", session.get("sello_page", 1)) or 1)
    firma_page = int(form.get("firma_page", session.get("firma_page", 1)) or 1)

    generados = []
    for id_txt in ids_seleccionados:
        idx = int(id_txt)
        if idx < 0 or idx >= len(session["rows"]):
            continue

        # Construir datos dinámicamente desde el formulario
        datos = {}
        for var in session["variables"]:
            valor_form = form.get(f"{var}_{idx}")
            datos[var] = str(valor_form) if valor_form else session["rows"][idx].get(var, "")

        # Usar primera variable para nombrar archivo
        primera_var = list(session["variables"])[0] if session["variables"] else "certificado"
        nombre_archivo = limpiar_nombre_archivo(datos.get(primera_var, "certificado"))
        ruta_docx = os.path.join("output", "certificados", f"{nombre_archivo}.docx")

        try:
            render_docx_desde_datos(
                session["plantilla_path"],
                datos=datos,
                ruta_docx=ruta_docx,
                no_valido=False,
                sello_path=None,
                firma_path=None,
                sello_x=0.0,
                sello_y=0.0,
                firma_x=0.0,
                firma_y=0.0,
                incluir_imagenes_docx=False,
                sello_width_in=DEFAULT_SELLO_WIDTH_IN,
                firma_width_in=DEFAULT_FIRMA_WIDTH_IN,
            )
            ruta_pdf_base = convertir_docx_a_pdf(ruta_docx, os.path.join("output", "certificados"))
            ruta_pdf = ruta_pdf_base.replace(".pdf", "_overlay.pdf")
            overlays = []
            
            # Usar posiciones individuales si existen, si no usar globales
            rows_overrides = session.get("rows_overrides", {})
            imagenes_source = session.get("imagenes", [])
            if str(idx) in rows_overrides:
                imagenes_source = rows_overrides[str(idx)]
            
            for imagen in imagenes_source:
                overlays.append(
                    {
                        "path": imagen.get("path"),
                        "x_in": imagen.get("x", 0.0),
                        "y_in": imagen.get("y", 0.0),
                        "width_in": imagen.get("width", DEFAULT_SELLO_WIDTH_IN),
                        "page": imagen.get("page", 1),
                    }
                )
            overlay_imagenes_en_pdf(
                ruta_pdf_entrada=ruta_pdf_base,
                ruta_pdf_salida=ruta_pdf,
                overlays=overlays,
            )
        except (RuntimeError, Exception) as exc:
            raise HTTPException(status_code=500, detail=str(exc))

        generados.append(
            {
                "nombre": nombre_archivo,
                "pdf": "/output/certificados/" + os.path.basename(ruta_pdf),
                "docx": "/output/certificados/" + os.path.basename(ruta_docx),
                "row_index": idx,
            }
        )

    if ticket_id:
        ticket = get_ticket(ticket_id)
        if not ticket:
            raise HTTPException(status_code=404, detail="Ticket no encontrado")
        if ticket.get("created_by") != user.get("username"):
            raise HTTPException(status_code=403, detail="No tienes permiso para editar este ticket")
        ticket.setdefault("generated", [])
        ticket["generated"].extend(generados)
        ticket["imagenes"] = session.get("imagenes", ticket.get("imagenes", []))
        ticket["updated_at"] = datetime.datetime.utcnow().isoformat()
        tickets = load_tickets()
        for idx_ticket, stored_ticket in enumerate(tickets):
            if stored_ticket.get("id") == ticket_id:
                tickets[idx_ticket] = ticket
                break
        save_tickets(tickets)
    else:
        new_ticket_id = uuid.uuid4().hex
        ticket_folder = save_ticket_assets(session, new_ticket_id)
        ticket_entry = {
            "id": new_ticket_id,
            "created_at": datetime.datetime.utcnow().isoformat(),
            "created_by": user.get("username"),
            "plantilla_nombre": session.get("plantilla_nombre"),
            "imagenes": [
                {
                    "filename": imagen.get("filename"),
                    "original_name": imagen.get("original_name"),
                    "x": imagen.get("x", 0.0),
                    "y": imagen.get("y", 0.0),
                    "width": imagen.get("width", DEFAULT_SELLO_WIDTH_IN),
                    "page": imagen.get("page", 1),
                }
                for imagen in session.get("imagenes", [])
            ],
            "page_width_in": session.get("page_width_in", 8.27),
            "page_height_in": session.get("page_height_in", 11.69),
            "variables": sorted(session.get("variables", [])),
            "rows": session.get("rows", []),
            "generated": generados,
        }
        tickets = load_tickets()
        tickets.append(ticket_entry)
        save_tickets(tickets)

    return templates.TemplateResponse(
        request,
        "resultado.html",
        {
            "request": request,
            "user": user,
            "generados": generados,
        },
    )
